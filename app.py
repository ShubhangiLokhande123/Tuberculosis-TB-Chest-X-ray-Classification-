"""
app.py — Standalone Gradio web app for TB Chest X-Ray Classification.
Deploy on Hugging Face Spaces (sdk: gradio) or run locally with `python app.py`.

Requirements: torch torchvision gradio python-docx opencv-python-headless Pillow numpy
Model weights: models/efficientnet_b0_tb.pth
"""

import io
import os
import numpy as np
import cv2
import torch
import torch.nn as nn
import torch.nn.functional as F
from torchvision import transforms
from torchvision.models import efficientnet_b0
from PIL import Image
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import gradio as gr

# ── Constants ─────────────────────────────────────────────────────────────────
CLASS_NAMES   = ['Normal', 'Tuberculosis']
WEIGHTS_PATH  = 'models/efficientnet_b0_tb.pth'
DEVICE        = torch.device('cuda' if torch.cuda.is_available() else 'cpu')

_infer_transforms = transforms.Compose([
    transforms.Resize((224, 224)),
    transforms.ToTensor(),
    transforms.Normalize(mean=[0.485, 0.456, 0.406],
                         std=[0.229, 0.224, 0.225])
])

# ── Model ─────────────────────────────────────────────────────────────────────
def build_model(weights_path: str, device):
    """Rebuild EfficientNet-B0 with the custom head and load saved weights."""
    model = efficientnet_b0(weights=None)
    in_features = model.classifier[1].in_features  # 1280
    model.classifier = nn.Sequential(
        nn.Dropout(0.5),
        nn.Linear(in_features, 256),
        nn.ReLU(inplace=True),
        nn.Dropout(0.3),
        nn.Linear(256, 2)
    )
    state = torch.load(weights_path, map_location=device, weights_only=True)
    model.load_state_dict(state)
    model.to(device)
    model.eval()
    return model


try:
    model = build_model(WEIGHTS_PATH, DEVICE)
    model_status = f"✅ Model loaded from '{WEIGHTS_PATH}' on {DEVICE}"
except FileNotFoundError:
    model = None
    model_status = f"⚠️ Weight file not found at '{WEIGHTS_PATH}'."
print(model_status)

# ── Inference utilities ───────────────────────────────────────────────────────
# Decision threshold tuned for best balanced accuracy on the target dataset.
# Raises specificity (Normal recall) from 26.7% to 82.7% vs default 0.50.
THRESHOLD = 0.945


def preprocess_image(pil_image: Image.Image):
    img_rgb     = pil_image.convert('RGB')
    img_resized = img_rgb.resize((224, 224), Image.LANCZOS)  # for display overlay
    rgb_float   = np.array(img_resized, dtype=np.float32) / 255.0
    tensor      = _infer_transforms(img_rgb).unsqueeze(0)
    return tensor, rgb_float


def predict(tensor):
    with torch.no_grad():
        probs = F.softmax(model(tensor.to(DEVICE)), dim=1)[0]
    normal_prob = float(probs[0])
    tb_prob     = float(probs[1])
    label = 'Tuberculosis' if tb_prob >= THRESHOLD else 'Normal'
    return label, normal_prob, tb_prob


def generate_gradcam(tensor, rgb_float):
    """Manual Grad-CAM via PyTorch forward/backward hooks on model.features[-1]."""
    tensor = tensor.to(DEVICE)
    gradients, activations = [], []

    def _fwd(module, inp, out):
        activations.append(out)

    def _bwd(module, grad_in, grad_out):
        gradients.append(grad_out[0])

    fh = model.features[-1].register_forward_hook(_fwd)
    bh = model.features[-1].register_full_backward_hook(_bwd)
    try:
        logits = model(tensor)
        model.zero_grad()
        logits[0, logits.argmax(dim=1).item()].backward()
    finally:
        fh.remove()
        bh.remove()

    weights = gradients[0].mean(dim=(2, 3), keepdim=True)
    cam     = F.relu((weights * activations[0]).sum(dim=1).squeeze())
    cam     = cam.detach().cpu().numpy()
    cam     = (cam - cam.min()) / (cam.max() + 1e-8)

    cam_r   = cv2.resize(cam, (224, 224))
    heatmap = cv2.applyColorMap(np.uint8(255 * cam_r), cv2.COLORMAP_JET)
    heatmap = cv2.cvtColor(heatmap, cv2.COLOR_BGR2RGB).astype(np.float32) / 255.0
    overlay = np.uint8(255 * np.clip(0.5 * rgb_float + 0.5 * heatmap, 0, 1))
    return Image.fromarray(overlay)

# ── DOCX report ───────────────────────────────────────────────────────────────
def _pil_to_stream(pil_img):
    buf = io.BytesIO()
    pil_img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def _add_heading(doc, text, level=2, color=RGBColor(0x1A, 0x23, 0x5E)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def create_docx_report(original_pil, heatmap_pil, class_name, normal_prob, tb_prob, timestamp):
    os.makedirs('temp', exist_ok=True)
    safe_ts  = timestamp.replace(':', '-').replace(' ', '_')
    out_path = os.path.join('temp', f'TB_Report_{safe_ts}.docx')

    is_tb        = class_name == 'Tuberculosis'
    result_color = RGBColor(0xD3, 0x2F, 0x2F) if is_tb else RGBColor(0x2E, 0x7D, 0x32)
    brand_color  = RGBColor(0x1A, 0x23, 0x5E)
    doc          = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title_para.add_run('TB Chest X-Ray Analysis Report')
    t.bold = True; t.font.size = Pt(20); t.font.color.rgb = brand_color

    ts_para = doc.add_paragraph()
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts_r = ts_para.add_run(f'Analysis Date: {timestamp}')
    ts_r.font.size = Pt(10); ts_r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    # Section 1: Classification Result
    _add_heading(doc, '1. Classification Result')
    res_para = doc.add_paragraph()
    res_run = res_para.add_run(f'Prediction:  {class_name}')
    res_run.bold = True; res_run.font.size = Pt(13); res_run.font.color.rgb = result_color
    doc.add_paragraph()

    conf_tbl = doc.add_table(rows=3, cols=2)
    conf_tbl.style = 'Table Grid'
    for col_idx, hdr in enumerate(['Class', 'Confidence (%)']):
        cell = conf_tbl.rows[0].cells[col_idx]
        cell.text = hdr
        cell.paragraphs[0].runs[0].bold = True
    conf_tbl.rows[1].cells[0].text = 'Normal'
    conf_tbl.rows[1].cells[1].text = f'{normal_prob * 100:.2f}%'
    conf_tbl.rows[2].cells[0].text = 'Tuberculosis'
    conf_tbl.rows[2].cells[1].text = f'{tb_prob * 100:.2f}%'
    detected_row = conf_tbl.rows[2] if is_tb else conf_tbl.rows[1]
    for cell in detected_row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True; run.font.color.rgb = result_color
    doc.add_paragraph()

    # Section 2: X-Ray Images
    _add_heading(doc, '2. X-Ray Images')
    img_tbl = doc.add_table(rows=2, cols=2)
    img_tbl.style = 'Table Grid'
    for col_idx, label in enumerate(['Original Chest X-Ray', 'Grad-CAM Heatmap Overlay']):
        cell = img_tbl.rows[0].cells[col_idx]
        cell.text = label
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for col_idx, stream in enumerate([_pil_to_stream(original_pil), _pil_to_stream(heatmap_pil)]):
        cell = img_tbl.rows[1].cells[col_idx]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(stream, width=Inches(2.8))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

    # Section 3: Clinical Recommendation
    _add_heading(doc, '3. Clinical Recommendation')
    rec_text = (
        'ALERT: The AI model has detected radiological features consistent with '
        'Tuberculosis (TB) in this chest X-ray. Immediate consultation with a '
        'pulmonologist or infectious disease specialist is strongly recommended. '
        'Further confirmatory testing (sputum culture, NAAT, TST/IGRA) should be '
        'initiated as soon as possible.'
        if is_tb else
        'The AI model did not detect significant radiological features associated '
        'with Tuberculosis in this chest X-ray. The image appears consistent with a '
        'normal chest X-ray. Routine clinical follow-up is advised as indicated.'
    )
    rec_para = doc.add_paragraph()
    rec_run = rec_para.add_run(rec_text)
    rec_run.font.size = Pt(11); rec_run.font.color.rgb = result_color
    doc.add_paragraph()

    # Disclaimer
    disc_para = doc.add_paragraph()
    disc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_run = disc_para.add_run(
        'DISCLAIMER: This report is generated by an AI-assisted tool for research and '
        'educational purposes only. It does NOT constitute a medical diagnosis and '
        'must NOT replace professional clinical judgment. Always consult a qualified '
        'healthcare professional for diagnosis and treatment decisions.'
    )
    disc_run.font.size = Pt(8); disc_run.italic = True
    disc_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(out_path)
    return out_path

# ── Gradio analysis callback ──────────────────────────────────────────────────
def analyze(image: Image.Image):
    if model is None:
        return None, None, f"### ⚠️ Model Not Loaded\n\n{model_status}", None
    if image is None:
        return None, None, "### Please upload a chest X-ray image to begin analysis.", None

    tensor, rgb_float = preprocess_image(image)
    class_name, normal_prob, tb_prob = predict(tensor)
    heatmap_pil = generate_gradcam(tensor, rgb_float)
    ts          = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    docx_path   = create_docx_report(image, heatmap_pil, class_name, normal_prob, tb_prob, ts)

    icon  = '🔴' if class_name == 'Tuberculosis' else '🟢'
    alert = (
        '> ⚠️ **ALERT:** Radiological features consistent with TB were detected.  \n'
        '> Immediate consultation with a specialist is strongly recommended.'
        if class_name == 'Tuberculosis' else
        '> ✅ No significant TB features detected.  \n'
        '> Routine clinical follow-up is advised as clinically indicated.'
    )
    report_md = f"""
## {icon} Diagnosis: **{class_name}**

| Class | Confidence |
|---|---|
| Normal | {normal_prob * 100:.2f}% |
| Tuberculosis | {tb_prob * 100:.2f}% |

**Analysis Timestamp:** {ts}

{alert}

---
*This report is generated by an AI-assisted tool for research and educational purposes only.  
It does **not** constitute a medical diagnosis and must not replace professional clinical judgment.*
"""
    return image, heatmap_pil, report_md, docx_path

# ── Gradio UI ─────────────────────────────────────────────────────────────────
with gr.Blocks(
    title='TB Chest X-Ray Classifier',
    theme=gr.themes.Soft(),
    css='.gradio-container { max-width: 1100px; margin: auto; }'
) as demo:

    gr.Markdown(
        """
        # 🫁 TB Chest X-Ray Analysis System
        ### EfficientNet-B0 · Grad-CAM Explainability · DOCX Report
        Upload a posterior–anterior (PA) chest X-ray to receive an instant AI-powered
        classification, a Grad-CAM saliency heatmap, and a downloadable diagnostic report.
        """
    )

    with gr.Row():
        with gr.Column(scale=1, min_width=280):
            input_img   = gr.Image(type='pil', label='Upload Chest X-Ray',
                                   height=300, image_mode='RGB')
            analyze_btn = gr.Button('🔍  Analyze', variant='primary', size='lg')

        with gr.Column(scale=2):
            with gr.Row():
                orig_out = gr.Image(label='Original X-Ray',   height=290, interactive=False)
                heat_out = gr.Image(label='Grad-CAM Heatmap', height=290, interactive=False)

    report_out   = gr.Markdown(label='Diagnostic Report')
    download_out = gr.File(label='📥 Download DOCX Report', interactive=False)

    analyze_btn.click(
        fn=analyze,
        inputs=[input_img],
        outputs=[orig_out, heat_out, report_out, download_out]
    )

    gr.Markdown(
        """
        ---
        **Disclaimer:** This tool is intended for research and educational use only.
        Predictions are not a substitute for professional medical diagnosis.
        """
    )

if __name__ == '__main__':
    demo.launch(server_name='0.0.0.0', server_port=7860)
